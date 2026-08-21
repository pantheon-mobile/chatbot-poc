import io
import time

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


def _iter_block_items(parent):
    parent_element = parent.element.body if isinstance(parent, DocumentObject) else parent._tc
    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _paragraph_text(paragraph: Paragraph) -> str:
    pieces = []
    for child in paragraph._p.iterchildren():
        if child.tag == qn("w:hyperlink"):
            relation_id = child.get(qn("r:id"))
            text = "".join(node.text or "" for node in child.iter(qn("w:t")))
            relation = paragraph.part.rels.get(relation_id) if relation_id else None
            pieces.append(f"[{text}]({relation.target_ref})" if relation and text else text)
        elif child.tag == qn("w:r"):
            pieces.append("".join(node.text or "" for node in child.iter(qn("w:t"))))
            if child.find(qn("w:br")) is not None:
                pieces.append("\n")
    return "".join(pieces).strip()


def _page_break_present(paragraph: Paragraph) -> bool:
    for br in paragraph._p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return paragraph._p.find(".//" + qn("w:lastRenderedPageBreak")) is not None


def _paragraph_kind(paragraph: Paragraph) -> tuple[str, int]:
    style_name = (paragraph.style.name if paragraph.style else "") or ""
    if style_name.lower().startswith("heading"):
        digits = "".join(character for character in style_name if character.isdigit())
        return "heading", int(digits or 1)
    properties = paragraph._p.pPr
    numbered = properties is not None and properties.numPr is not None
    if numbered:
        if "bullet" in style_name.lower() or "箇条書き" in style_name:
            return "bullet", 0
        return "numbered", 0
    if "bullet" in style_name.lower() or "箇条書き" in style_name:
        return "bullet", 0
    if "number" in style_name.lower() or "番号" in style_name:
        return "numbered", 0
    return "paragraph", 0


def _table_rows(table: Table) -> list[list[str]]:
    rows = []
    seen_cells = set()
    for row in table.rows:
        values = []
        for cell in row.cells:
            cell_identity = id(cell._tc)
            if cell_identity in seen_cells:
                values.append("")
                continue
            seen_cells.add(cell_identity)
            values.append("\n".join(p.text for p in cell.paragraphs if p.text).strip())
        rows.append(values)
    return rows


def _markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    escaped = [[value.replace("|", "\\|").replace("\n", "<br>") for value in row]
               for row in normalized]
    return "\n".join([
        "| " + " | ".join(escaped[0]) + " |",
        "| " + " | ".join(["---"] * width) + " |",
        *("| " + " | ".join(row) + " |" for row in escaped[1:])
    ])


def _header_footer(document) -> tuple[list[str], list[str]]:
    headers, footers = [], []
    for section in document.sections:
        header_text = "\n".join(p.text for p in section.header.paragraphs if p.text.strip()).strip()
        footer_text = "\n".join(p.text for p in section.footer.paragraphs if p.text.strip()).strip()
        if header_text and header_text not in headers:
            headers.append(header_text)
        if footer_text and footer_text not in footers:
            footers.append(footer_text)
    return headers, footers


def parse_docx_document(docx_bytes: bytes, original_name: str) -> dict:
    load_started = time.perf_counter()
    document = Document(io.BytesIO(docx_bytes))
    docx_load_ms = round((time.perf_counter() - load_started) * 1000, 1)
    headers, footers = _header_footer(document)
    title = (document.core_properties.title or "").strip()
    blocks, paragraph_count, heading_count, table_count = [], 0, 0, 0
    paragraph_parse_ms, table_parse_ms = 0.0, 0.0
    for item in _iter_block_items(document):
        if isinstance(item, Paragraph):
            started = time.perf_counter()
            text = _paragraph_text(item)
            kind, level = _paragraph_kind(item)
            page_break = _page_break_present(item)
            if text or page_break:
                blocks.append({"type": kind, "level": level, "text": text, "page_break": page_break})
            if text:
                paragraph_count += 1
                if kind == "heading":
                    heading_count += 1
                    if not title and level == 1:
                        title = text
            paragraph_parse_ms += (time.perf_counter() - started) * 1000
        else:
            started = time.perf_counter()
            blocks.append({"type": "table", "rows": _table_rows(item)})
            table_count += 1
            table_parse_ms += (time.perf_counter() - started) * 1000

    txt_started = time.perf_counter()
    txt_parts = []
    if headers:
        txt_parts.append("--- Header ---\n" + "\n".join(headers))
    for block in blocks:
        if block["type"] == "heading":
            txt_parts.append(f"--- Heading {block['level']} ---\n{block['text']}")
        elif block["type"] == "bullet":
            txt_parts.append(f"--- Bullet List ---\n- {block['text']}")
        elif block["type"] == "numbered":
            txt_parts.append(f"--- Numbered List ---\n1. {block['text']}")
        elif block["type"] == "table":
            txt_parts.append("--- Table ---\n" + "\n".join(" | ".join(row) for row in block["rows"]))
        elif block["text"]:
            txt_parts.append(block["text"])
        if block.get("page_break"):
            txt_parts.append("--- Page Break ---")
    if footers:
        txt_parts.append("--- Footer ---\n" + "\n".join(footers))
    txt_text = "\n\n".join(txt_parts).strip()
    txt_generation_ms = round((time.perf_counter() - txt_started) * 1000, 1)

    markdown_started = time.perf_counter()
    markdown_parts = []
    if headers:
        markdown_parts.append("<!-- Header -->\n\n" + "\n\n".join(headers))
    for block in blocks:
        if block["type"] == "heading":
            markdown_parts.append(f"{'#' * min(block['level'], 6)} {block['text']}")
        elif block["type"] == "bullet":
            markdown_parts.append(f"- {block['text']}")
        elif block["type"] == "numbered":
            markdown_parts.append(f"1. {block['text']}")
        elif block["type"] == "table":
            markdown_parts.append(_markdown_table(block["rows"]))
        elif block["text"]:
            markdown_parts.append(block["text"])
        if block.get("page_break"):
            markdown_parts.append("<!-- Page Break -->\n\n---")
    if footers:
        markdown_parts.append("<!-- Footer -->\n\n" + "\n\n".join(footers))
    markdown_text = "\n\n".join(part for part in markdown_parts if part).strip()
    markdown_generation_ms = round((time.perf_counter() - markdown_started) * 1000, 1)
    return {
        "title": title, "paragraph_count": paragraph_count, "heading_count": heading_count,
        "table_count": table_count, "header_present": bool(headers), "footer_present": bool(footers),
        "docx_load_ms": docx_load_ms, "paragraph_parse_ms": round(paragraph_parse_ms, 1),
        "table_parse_ms": round(table_parse_ms, 1), "txt_generation_ms": txt_generation_ms,
        "markdown_generation_ms": markdown_generation_ms, "txt_text": txt_text,
        "markdown_text": markdown_text, "generated_file_count": 3,
        "generated_total_bytes": len(docx_bytes) + len(txt_text.encode("utf-8")) + len(markdown_text.encode("utf-8")),
        "unsupported_elements": "footnotes, comments, complex_text_boxes"
    }
