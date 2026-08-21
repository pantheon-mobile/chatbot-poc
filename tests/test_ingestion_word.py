from io import BytesIO

from docx import Document
from docx.enum.text import WD_BREAK

from ingestion_word import parse_docx_document


def _docx_bytes() -> bytes:
    document = Document()
    document.core_properties.title = "奨学金案内"
    document.sections[0].header.paragraphs[0].text = "大学ヘッダー"
    document.sections[0].footer.paragraphs[0].text = "1ページ"
    document.add_heading("対象者", level=1)
    document.add_paragraph("学生", style="List Bullet")
    document.add_paragraph("申請する", style="List Number")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "区分"
    table.cell(0, 1).text = "金額"
    table.cell(1, 0).text = "第1区分"
    table.cell(1, 1).text = "75,800円"
    paragraph = document.add_paragraph("次ページ")
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_structure_is_preserved_in_txt_and_markdown():
    result = parse_docx_document(_docx_bytes(), "sample.docx")

    assert result["title"] == "奨学金案内"
    assert result["heading_count"] == 1
    assert result["table_count"] == 1
    assert result["header_present"] is True
    assert result["footer_present"] is True
    assert "--- Heading 1 ---\n対象者" in result["txt_text"]
    assert "--- Table ---" in result["txt_text"]
    assert "# 対象者" in result["markdown_text"]
    assert "| 区分 | 金額 |" in result["markdown_text"]
    assert "<!-- Page Break -->" in result["markdown_text"]


def test_docx_generation_metrics_are_recorded():
    result = parse_docx_document(_docx_bytes(), "sample.docx")

    assert result["generated_file_count"] == 3
    assert result["generated_total_bytes"] > len(_docx_bytes())
    assert result["docx_load_ms"] >= 0
    assert result["paragraph_parse_ms"] >= 0
    assert result["table_parse_ms"] >= 0
